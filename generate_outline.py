import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


def set_cell_background(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def add_bottom_border(paragraph, color="1B365D", sz="6", space="4"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)


NAVY = RGBColor(0x1B, 0x36, 0x5D)
TEAL = RGBColor(0x4A, 0x77, 0x7A)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x55, 0x55, 0x55)
FONT = "Calibri"


def make_run(paragraph, text, size=11, bold=False, italic=False, color=DARK_GRAY):
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def add_h1(doc, text):
    p = doc.add_paragraph()
    make_run(p, text, size=16, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    add_bottom_border(p)


def add_h2(doc, text):
    p = doc.add_paragraph()
    make_run(p, text, size=13, bold=True, color=TEAL)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def add_h3(doc, text):
    p = doc.add_paragraph()
    make_run(p, text, size=11, bold=True, italic=True, color=MID_GRAY)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        make_run(p, bold_prefix, bold=True)
    make_run(p, text)


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        make_run(p, bold_prefix, bold=True)
    make_run(p, text)
    p.paragraph_format.space_after = Pt(6)


def add_callout(doc, text):
    p = doc.add_paragraph()
    make_run(p, text, italic=True, color=TEAL)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(18)


# ── Build Document ──────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

# Default style
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style.font.color.rgb = DARK_GRAY
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

# ── Title ──
p = doc.add_paragraph()
make_run(p, "OUTLINE BÁO CÁO KẾT QUẢ DỰ ÁN", size=22, bold=True, color=NAVY)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(4)
p2 = doc.add_paragraph()
make_run(p2, "Hệ thống Cảnh báo Giao dịch Bất thường — Anomaly Alert System", size=14, color=TEAL)
p2.paragraph_format.space_after = Pt(20)

# ── Phần I ──
add_h1(doc, "Phần I: Giới thiệu")

add_h2(doc, "1.1. Bối cảnh & Vấn đề cần giải quyết")
add_bullet(doc, "Thực trạng gian lận tài chính qua kênh ngân hàng số tại Việt Nam (tăng 65%/năm, thiệt hại >8.000 tỷ VND)")
add_bullet(doc, "Các hình thức gian lận chính: chiếm đoạt tài khoản (ATO), tài khoản trung chuyển (Money Mule), chia nhỏ giao dịch (Structuring)")
add_bullet(doc, "Hệ thống cảnh báo truyền thống (rule-based) gây báo động giả >95%, làm giảm trải nghiệm khách hàng")

add_h2(doc, "1.2. Mục tiêu hệ thống")
add_bullet(doc, "Phát hiện giao dịch nghi vấn gian lận với tỷ lệ cảnh báo chính xác")
add_bullet(doc, "Giảm báo động giả — không khóa nhầm tài khoản khách hàng hợp lệ")
add_bullet(doc, "Tự động giải thích lý do cảnh báo bằng ngôn ngữ tự nhiên cho nhân viên Compliance")
add_bullet(doc, "Đề xuất bước xác thực bổ sung tối thiểu thay vì chặn giao dịch")
add_bullet(doc, "Bảo mật không ma sát (Frictionless Security) — giữ chân khách hàng, tăng uy tín ngân hàng")

add_h2(doc, "1.3. Phạm vi dự án")
add_bullet(doc, "Lựa chọn Hướng 1: Fraud & Anomaly Detection")
add_bullet(doc, "Sơ đồ tổng quan kiến trúc hệ thống (5 module xử lý tuần tự)")

add_h2(doc, "1.4. Dữ liệu sử dụng")
add_bullet(doc, "6 bảng dữ liệu ngân hàng: Khách hàng, Giao dịch, Hoạt động số, Tiền gửi, Cho vay, Thẻ")
add_bullet(doc, "Quy mô: ~20 triệu bản ghi")

# ── Phần II ──
add_h1(doc, "Phần II: Tổng quan Dữ liệu (EDA)")

add_h2(doc, "2.1. Quy mô & cấu trúc dữ liệu")
add_bullet(doc, "Tổng số khách hàng, giao dịch, hoạt động")
add_bullet(doc, "Phân bố giao dịch theo thời gian (ngày, giờ, ngày trong tuần)")

add_h2(doc, "2.2. Phân bố giao dịch")
add_bullet(doc, "Phân bố số tiền giao dịch (phần lớn nhỏ lẻ, thiểu số giá trị rất lớn)")
add_bullet(doc, "Tỷ trọng theo loại giao dịch (Transfer Outside, Within, eWallet...)")
add_bullet(doc, "Tỷ trọng theo thiết bị (iOS, Android, Web)")

add_h2(doc, "2.3. Phân tích hành vi theo nhóm khách hàng")
add_bullet(doc, "Khác biệt hành vi giao dịch giữa các nhóm tuổi, nghề nghiệp, thời gian sử dụng dịch vụ")
add_bullet(doc, "Nhận diện các nhóm có rủi ro cao hơn (sinh viên, tài khoản mới, tài khoản ngủ đông)")

add_h2(doc, "2.4. Phân tích hoạt động bảo mật")
add_bullet(doc, "Tần suất đổi mật khẩu/PIN và mối liên hệ với giao dịch sau đó")
add_bullet(doc, "Các insight định hướng thiết kế tín hiệu cảnh báo")

# ── Phần III ──
add_h1(doc, "Phần III: Hệ thống Thu thập & Xử lý Tín hiệu Rủi ro")

add_h2(doc, "3.1. Thu thập dữ liệu giao dịch thời gian thực")
add_bullet(doc, "Nối và làm giàu dữ liệu từ 6 nguồn")
add_bullet(doc, "Tính toán tổng tiền & tần suất giao dịch theo 6 khung thời gian (1 giờ → 30 ngày)")
add_bullet(doc, "Số dư trung bình & mức chi tiêu lịch sử của từng khách hàng")

add_h2(doc, "3.2. Các tín hiệu rủi ro được hệ thống theo dõi")
add_bullet(doc, "Giao dịch lớn bất thường so với lịch sử cá nhân? So với số dư? So với 30 ngày gần nhất?", bold_prefix="Tín hiệu số tiền: ")
add_bullet(doc, "Tần suất và tổng tiền tập trung bất thường trong 1 giờ / 24 giờ gần nhất?", bold_prefix="Tín hiệu tốc độ: ")
add_bullet(doc, "Tài khoản im lặng lâu rồi đột ngột giao dịch lớn? Giao dịch đêm bất thường so với thói quen?", bold_prefix="Tín hiệu thời gian: ")
add_bullet(doc, "Giao dịch xảy ra ngay sau khi đổi mật khẩu/PIN? Phương thức đăng nhập khác thường?", bold_prefix="Tín hiệu bảo mật: ")
add_bullet(doc, "Phân bố số tiền vi phạm Định luật Benford (dấu hiệu chia nhỏ có chủ đích)?", bold_prefix="Tín hiệu thống kê: ")
add_bullet(doc, "Chuỗi thao tác trên app bất thường so với hàng triệu phiên giao dịch bình thường?", bold_prefix="Tín hiệu hành vi: ")
add_bullet(doc, "Nhóm tuổi, nghề nghiệp, loại giao dịch — để cá nhân hóa ngưỡng rủi ro thay vì áp dụng một ngưỡng cứng cho tất cả", bold_prefix="Bối cảnh khách hàng: ")

add_h2(doc, "3.3. Tổng kết")
add_para(doc, "~46 tín hiệu rủi ro được tổng hợp cho mỗi giao dịch.")

# ── Phần IV ──
add_h1(doc, "Phần IV: Cơ chế Huấn luyện Mô hình")

add_h2(doc, "4.1. Thách thức: Không có dữ liệu gian lận được gắn nhãn sẵn")
add_bullet(doc, "Trong thực tế ngân hàng, không tồn tại nhãn \"giao dịch gian lận\" xác nhận → hệ thống phải tự học từ dữ liệu chưa gắn nhãn")

add_h2(doc, "4.2. Quy trình tự tạo nhãn và lọc nhiễu (3 bước)")
add_bullet(doc, "Dùng thuật toán phát hiện bất thường (Isolation Forest) để đánh dấu ~3% giao dịch nghi vấn nhất", bold_prefix="Bước 1 — Sàng lọc ban đầu: ")
add_bullet(doc, "\"Cài gián điệp\" vào nhóm bình thường để kiểm tra độ tin cậy → tách ra nhóm giao dịch bình thường đáng tin", bold_prefix="Bước 2 — Xác minh chéo: ")
add_bullet(doc, "Kiểm chứng chéo 5 lần trên nhóm chưa rõ ràng, loại bỏ 10% mẫu có xác suất sai nhãn cao nhất", bold_prefix="Bước 3 — Loại bỏ nhiễu: ")

add_h2(doc, "4.3. Huấn luyện bộ phân loại")
add_bullet(doc, "Thuật toán XGBoost với cấu hình chống quá khớp (overfitting)")
add_bullet(doc, "Hiệu chỉnh xác suất đầu ra (Elkan-Noto) cho ra điểm rủi ro chính xác")

add_h2(doc, "4.4. Ngưỡng quyết định")
add_bullet(doc, "Top 3% giao dịch có điểm rủi ro cao nhất được gắn cờ cảnh báo")

# ── Phần V ──
add_h1(doc, "Phần V: Cơ chế Xử lý Giao dịch 3 Tầng (Tiered Inference)")
add_callout(doc, "Đây là kiến trúc cốt lõi của hệ thống — phân luồng thông minh để vừa nhanh, vừa chính xác, vừa giải thích được.")

add_h2(doc, "5.1. Tổng quan: Tại sao cần 3 tầng?")
add_bullet(doc, "Không phải mọi giao dịch đều cần qua mô hình AI → lãng phí tài nguyên, tăng độ trễ")
add_bullet(doc, "Hệ thống 3 tầng: giao dịch rõ ràng an toàn → bỏ qua ngay; giao dịch mơ hồ → AI đánh giá; giao dịch bị gắn cờ → giải thích chi tiết")

add_h2(doc, "5.2. Tầng 1 — Bỏ qua Nhanh (Rule-Based Bypass)")
add_bullet(doc, "Giao dịch nhỏ (<500.000 VND) + tần suất thấp + hành vi bình thường → xử lý tức thì, không qua AI")
add_bullet(doc, "~95% giao dịch được xử lý trong <1ms → không gây trễ cho khách hàng", bold_prefix="Kết quả: ")

add_h2(doc, "5.3. Tầng 2 — Đánh giá bằng AI (ML Classification)")
add_bullet(doc, "~5% giao dịch còn lại được chấm điểm rủi ro bằng mô hình XGBoost")
add_bullet(doc, "Điểm rủi ro (0-100%) + quyết định gắn cờ/không gắn cờ", bold_prefix="Output: ")

add_h2(doc, "5.4. Tầng 3 — Giải thích Tự động (xAI Engine)")
add_bullet(doc, "Chỉ kích hoạt cho các giao dịch bị gắn cờ bất thường")

add_h3(doc, "5.4.1. Lý do cảnh báo bằng ngôn ngữ tự nhiên")
add_bullet(doc, "Hệ thống tự sinh câu giải thích, ví dụ: \"Số tiền gấp 15 lần mức giao dịch trung bình của khách hàng\"")

add_h3(doc, "5.4.2. Phân tích tổ hợp rủi ro")
add_bullet(doc, "Phát hiện các cặp tín hiệu kết hợp nguy hiểm, ví dụ: \"Số tiền lớn bất thường + tài khoản ngủ đông lâu\"")

add_h3(doc, "5.4.3. Đề xuất hành động khắc phục (Counterfactual Recourse)")
add_bullet(doc, "Thay vì chặn cứng, hệ thống đề xuất: \"Giảm số tiền xuống 45 triệu\" hoặc \"Xác minh FaceID để tiếp tục\"")
add_bullet(doc, "Tự động tính toán mức thay đổi tối thiểu để giao dịch thoát ngưỡng cảnh báo")

# ── Phần VI ──
add_h1(doc, "Phần VI: Khả năng Tự cập nhật (Continuous Learning)")

add_h2(doc, "6.1. Vấn đề: Hành vi giao dịch thay đổi theo thời gian")
add_bullet(doc, "Khách hàng thay đổi thói quen chi tiêu, mùa lễ Tết tăng giao dịch → mô hình cần thích ứng mà không \"quên\" kiến thức cũ")

add_h2(doc, "6.2. Giải pháp: Cơ chế bảo vệ trí nhớ (EWC)")
add_bullet(doc, "Mô hình tự cập nhật khi có dữ liệu mới nhưng vẫn giữ được khả năng phát hiện các mẫu gian lận đã học")

add_h2(doc, "6.3. Kết quả kiểm chứng")
add_bullet(doc, "So sánh mô hình Có/Không có cơ chế bảo vệ khi dữ liệu thay đổi mạnh (x5 số tiền)")
add_bullet(doc, "Mô hình được bảo vệ duy trì hiệu quả phát hiện trên dữ liệu gốc")

# ── Phần VII ──
add_h1(doc, "Phần VII: Kết quả Hệ thống")

add_h2(doc, "7.1. Hiệu quả phát hiện")
add_bullet(doc, "Tổng giao dịch đánh giá: 5.000")
add_bullet(doc, "Giao dịch bị gắn cờ: 121 (2,42%) — thấp hơn đáng kể so với 5-10% báo động giả của hệ thống cũ")
add_bullet(doc, "66 khách hàng bị ảnh hưởng, cao nhất 9 cảnh báo/1 khách hàng")
add_bullet(doc, "Giá trị trung bình giao dịch nghi vấn: 42 triệu VND (trung vị: 21 triệu VND)")

add_h2(doc, "7.2. Phân bố rủi ro — Hệ thống nhắm đúng mục tiêu")
add_bullet(doc, "57,85% cảnh báo là chuyển tiền ra ngoài ngân hàng — đúng nhóm rủi ro cao nhất, khó thu hồi nhất", bold_prefix="Kết quả: ")
add_bullet(doc, "Thiết bị: iOS (50,4%), Android (47,9%), Web (1,7%)")

add_h2(doc, "7.3. Các tín hiệu rủi ro quan trọng nhất (Top SHAP Contributors)")
add_bullet(doc, "Bảng top 10 tín hiệu xuất hiện nhiều nhất trong 121 cảnh báo")
add_bullet(doc, "\"Giao dịch chiếm tỷ trọng lớn so với số dư\" xuất hiện trong 80% cảnh báo — khớp với mô hình tài khoản trung chuyển (Money Mule)", bold_prefix="Phân tích: ")

add_h2(doc, "7.4. Các tổ hợp rủi ro nguy hiểm nhất")
add_bullet(doc, "Top 5 cặp tín hiệu kết hợp và ý nghĩa thực tế")

add_h2(doc, "7.5. Ví dụ Đề xuất Khắc phục (Counterfactual Recourse)")
add_bullet(doc, "Case study cụ thể: giao dịch bị cờ → hệ thống đề xuất giảm số tiền hoặc xác minh sinh trắc học")

add_h2(doc, "7.6. Hiệu suất xử lý")
add_bullet(doc, "Tầng 1 (Rule Bypass): 0,83 ms — tức thì")
add_bullet(doc, "Tầng 2 (AI Classification): 39,86 ms — gần real-time")
add_bullet(doc, "Tầng 3 (Giải thích xAI): 8,86 giây — chỉ cho giao dịch bị cờ")
add_bullet(doc, "→ Đủ nhanh cho hệ thống cảnh báo thời gian thực")

# ── Phần VIII ──
add_h1(doc, "Phần VIII: Giá trị Kinh doanh & Đề xuất Triển khai")

add_h2(doc, "8.1. Bảo vệ uy tín ngân hàng")
add_bullet(doc, "Tỷ lệ cảnh báo 2,42% vs 5-10% báo động giả của hệ thống rule-based cũ")
add_bullet(doc, "Tập trung đúng vào giao dịch rủi ro cao (chuyển tiền ra ngoài chiếm 57,85% cảnh báo)")

add_h2(doc, "8.2. Nâng cao trải nghiệm khách hàng")
add_bullet(doc, "95% giao dịch xử lý tức thì không qua AI → khách hàng bình thường không bị ảnh hưởng")
add_bullet(doc, "Đề xuất xác thực bổ sung (FaceID/OTP) thay vì chặn cứng → giảm ma sát, giữ chân khách hàng")

add_h2(doc, "8.3. Tăng hiệu suất đội ngũ Compliance")
add_bullet(doc, "Mỗi cảnh báo đi kèm giải thích tự động bằng ngôn ngữ tự nhiên → Analyst duyệt nhanh gấp 3-5 lần")
add_bullet(doc, "Phân tích tổ hợp rủi ro cho biết chính xác sự kết hợp yếu tố nào gây ra cảnh báo → không cần phỏng đoán")

add_h2(doc, "8.4. Khả năng mở rộng & Hướng phát triển")
add_bullet(doc, "Tích hợp feedback loop từ Analyst để mô hình tự cải thiện")
add_bullet(doc, "Nâng cấp lên pipeline streaming thời gian thực (Kafka)")
add_bullet(doc, "Bổ sung phân tích mạng lưới giao dịch (Graph-based detection) để phát hiện vòng rửa tiền")

# ── Phần IX ──
add_h1(doc, "Phần IX: Tài liệu Tham khảo")

import shutil

output_path = "outline.docx"
temp_path = "outline_v2.docx"
doc.save(temp_path)

try:
    shutil.move(temp_path, output_path)
    print(f"Successfully generated {output_path}")
except PermissionError:
    print(f"WARNING: {output_path} is locked (likely open in Word). Saved as {temp_path} instead.")
    print(f"Close {output_path} in Word, then rename {temp_path} → {output_path}.")
